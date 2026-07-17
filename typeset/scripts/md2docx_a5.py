# -*- coding: utf-8 -*-
"""Markdown -> DOCX 변환기 (부크크 A5/국판 단행본용).

B5 교재는 hwpx 파이프라인(md2hwpx.py)을 쓰지만, A5 단행본(논문작성법 류)은
부크크 워드서식 docx 를 템플릿으로 python-docx 로 조판한다.
원본: U:\\2026\\02_논문POD집필\\build_bookk_a5_chapters.py (AI활용논문작성법 v3 출간 검증)

페이지: 15.4 x 21.6 cm (부크크 A5), 여백 좌우 2.3 / 상 2.8 / 하 3.3 / 제본 0.5 cm
글꼴: 본문 KoPubWorld바탕체 9.4pt(행간 1.45), 제목 KoPubWorld돋움체, 코드 Consolas

의존성: pip install markdown beautifulsoup4 lxml python-docx

사용:
  python md2docx_a5.py --out <출력.docx> 장01.md [장02.md ...]      # md 여러 개 = 한 문서로 연결
  python md2docx_a5.py --out-dir <폴더> 장01.md 장02.md ...          # md 하나당 docx 하나
  (--template 생략 시 부크크 기본 A5 워드서식 사용)
"""
from __future__ import annotations

import argparse
import re
from pathlib import Path

import markdown
from bs4 import BeautifulSoup, NavigableString
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# 기본 템플릿: 스킬 동봉본(assets/) 우선, 없으면 이 PC의 원본 경로로 폴백
_ASSET_TEMPLATE = Path(__file__).resolve().parent.parent / "assets" / "A5_bookk_template.docx"
_LOCAL_TEMPLATE = Path(r"U:\2026\02_논문POD집필\부크크서식\[워드서식]A5(국판)_부크크서식(기본).docx")
DEFAULT_TEMPLATE = str(_ASSET_TEMPLATE if _ASSET_TEMPLATE.exists() else _LOCAL_TEMPLATE)

BODY_FONT = "KoPubWorld바탕체"
HEAD_FONT = "KoPubWorld돋움체"
FALLBACK_FONT = "맑은 고딕"
MONO_FONT = "Consolas"


def clear_body_keep_section(doc: Document) -> None:
    body = doc._body._element
    sect_pr = body.sectPr
    for child in list(body):
        body.remove(child)
    if sect_pr is not None:
        body.append(sect_pr)


def set_run_font(run, *, east_asia=BODY_FONT, latin=FALLBACK_FONT, size=None,
                 bold=None, italic=None, color=None, mono=False) -> None:
    if mono:
        east_asia = MONO_FONT
        latin = MONO_FONT
    run.font.name = latin
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), east_asia)
    r_fonts.set(qn("w:ascii"), latin)
    r_fonts.set(qn("w:hAnsi"), latin)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=90, bottom=80, end=90) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def setup_bookk_page(doc: Document) -> None:
    """부크크 A5(국판) 페이지 설정."""
    for section in doc.sections:
        section.start_type = WD_SECTION_START.NEW_PAGE
        section.page_width = Cm(15.4)
        section.page_height = Cm(21.6)
        section.left_margin = Cm(2.3)
        section.right_margin = Cm(2.3)
        section.top_margin = Cm(2.8)
        section.bottom_margin = Cm(3.3)
        section.gutter = Cm(0.5)
        section.header_distance = Cm(1.2)
        section.footer_distance = Cm(1.4)


def setup_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = FALLBACK_FONT
    normal.font.size = Pt(9.4)
    normal._element.get_or_add_rPr().append(OxmlElement("w:rFonts"))
    r_fonts = normal._element.get_or_add_rPr().find(qn("w:rFonts"))
    r_fonts.set(qn("w:eastAsia"), BODY_FONT)
    r_fonts.set(qn("w:ascii"), FALLBACK_FONT)
    r_fonts.set(qn("w:hAnsi"), FALLBACK_FONT)
    normal.paragraph_format.line_spacing = 1.45
    normal.paragraph_format.space_after = Pt(3)


def add_page_number_footer(doc: Document) -> None:
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.text = ""
        run = p.add_run()
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = "PAGE"
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        run._r.append(fld_begin)
        run._r.append(instr)
        run._r.append(fld_end)
        set_run_font(run, east_asia=HEAD_FONT, size=8.5, color="777777")


def add_inline(paragraph, element, *, bold=False, italic=False, mono=False, size=None) -> None:
    for child in element.children:
        if isinstance(child, NavigableString):
            text = str(child)
            if not text:
                continue
            run = paragraph.add_run(text)
            set_run_font(run, size=size, bold=bold or None, italic=italic or None, mono=mono)
            continue
        name = child.name
        if name in ("strong", "b"):
            add_inline(paragraph, child, bold=True, italic=italic, mono=mono, size=size)
        elif name in ("em", "i"):
            add_inline(paragraph, child, bold=bold, italic=True, mono=mono, size=size)
        elif name == "code":
            add_inline(paragraph, child, bold=bold, italic=italic, mono=True, size=8.4)
        elif name == "br":
            paragraph.add_run().add_break()
        else:
            add_inline(paragraph, child, bold=bold, italic=italic, mono=mono, size=size)


def add_heading(doc: Document, text: str, level: int) -> None:
    text = re.sub(r"\s+", " ", text).strip()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.keep_with_next = True
    pf.first_line_indent = None

    if level == 1:
        pf.space_before = Pt(0)
        pf.space_after = Pt(14)
        run = p.add_run(text)
        set_run_font(run, east_asia=HEAD_FONT, size=19, bold=True, color="111111")
    elif level == 2 and text.startswith("제"):
        # "제N장 ..." 꼴은 장 표제로 크게, 가운데 정렬
        pf.space_before = Pt(18)
        pf.space_after = Pt(16)
        run = p.add_run(text)
        set_run_font(run, east_asia=HEAD_FONT, size=17, bold=True, color="111111")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 2:
        pf.space_before = Pt(16)
        pf.space_after = Pt(8)
        run = p.add_run(text)
        set_run_font(run, east_asia=HEAD_FONT, size=13.2, bold=True, color="1F4E79")
    elif level == 3:
        pf.space_before = Pt(12)
        pf.space_after = Pt(5)
        run = p.add_run(text)
        set_run_font(run, east_asia=HEAD_FONT, size=11.2, bold=True, color="222222")
    else:
        pf.space_before = Pt(8)
        pf.space_after = Pt(4)
        run = p.add_run(text)
        set_run_font(run, east_asia=HEAD_FONT, size=10.2, bold=True, color="333333")


def add_paragraph(doc: Document, element) -> None:
    if not element.get_text(strip=True):
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.line_spacing = 1.45
    pf.space_after = Pt(3.5)
    pf.first_line_indent = Cm(0.55)
    add_inline(p, element, size=9.4)


def add_blockquote(doc: Document, element) -> None:
    for child in element.children:
        if isinstance(child, NavigableString) and not str(child).strip():
            continue
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf = p.paragraph_format
        pf.left_indent = Cm(0.35)
        pf.right_indent = Cm(0.15)
        pf.first_line_indent = None
        pf.line_spacing = 1.35
        pf.space_before = Pt(2)
        pf.space_after = Pt(4)
        set_paragraph_shading(p, "F3F6F8")
        if isinstance(child, NavigableString):
            run = p.add_run(str(child).strip())
            set_run_font(run, size=8.8, italic=True)
        elif child.name == "p":
            add_inline(p, child, italic=False, size=8.8)
        elif child.name in ("ul", "ol"):
            for li in child.find_all("li", recursive=False):
                p.add_run("· ")
                add_inline(p, li, size=8.8)


def add_list(doc: Document, element, *, ordered=False, level=0) -> None:
    for index, li in enumerate(element.find_all("li", recursive=False), start=1):
        nested = li.find_all(["ul", "ol"], recursive=False)
        for node in nested:
            node.extract()
        style = "List Number" if ordered else "List Bullet"
        try:
            p = doc.add_paragraph(style=style)
        except KeyError:
            p = doc.add_paragraph()
            p.add_run(f"{index}. " if ordered else "- ")
        p.paragraph_format.left_indent = Cm(0.45 + 0.35 * level)
        p.paragraph_format.first_line_indent = Cm(-0.18)
        p.paragraph_format.line_spacing = 1.35
        p.paragraph_format.space_after = Pt(2.5)
        add_inline(p, li, size=9.1)
        for node in nested:
            add_list(doc, node, ordered=(node.name == "ol"), level=level + 1)


def add_table(doc: Document, element) -> None:
    rows = element.find_all("tr")
    if not rows:
        return
    cols = max(len(row.find_all(["td", "th"])) for row in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    try:
        table.style = "Table Grid"
    except KeyError:
        pass

    section = doc.sections[-1]
    usable = section.page_width - section.left_margin - section.right_margin - section.gutter
    total_width = int(usable * 0.94)
    col_width = int(total_width / max(cols, 1))
    for col in table.columns:
        for cell in col.cells:
            cell.width = col_width

    font_size = 8.0 if cols <= 4 else 6.6
    for row_idx, row in enumerate(rows):
        cells = row.find_all(["td", "th"])
        for col_idx, source_cell in enumerate(cells):
            target = table.rows[row_idx].cells[col_idx]
            target.text = ""
            target.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(target)
            p = target.paragraphs[0]
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_after = Pt(0)
            p.alignment = (WD_ALIGN_PARAGRAPH.CENTER
                           if cols > 4 or source_cell.get_text(strip=True) in {"●", "○"}
                           else WD_ALIGN_PARAGRAPH.LEFT)
            add_inline(p, source_cell, bold=(source_cell.name == "th" or row_idx == 0), size=font_size)
            if source_cell.name == "th" or row_idx == 0:
                set_cell_shading(target, "EAF0F5")
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_code_block(doc: Document, element) -> None:
    code = element.get_text().rstrip("\n")
    for line in code.splitlines() or [""]:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.35)
        p.paragraph_format.right_indent = Cm(0.15)
        p.paragraph_format.first_line_indent = None
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(0)
        set_paragraph_shading(p, "F5F5F5")
        run = p.add_run(line if line else " ")
        set_run_font(run, mono=True, size=7.4)


def add_hr(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "D9D9D9")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_image(doc: Document, img_el, base_dir: Path) -> None:
    src = img_el.get("src")
    if not src:
        return
    image_path = Path(src)
    if not image_path.is_absolute():
        image_path = base_dir / image_path
    if not image_path.exists() and not Path(src).is_absolute():
        image_path = Path.cwd() / src
    if not image_path.exists():
        return
    section = doc.sections[-1]
    usable_width = section.page_width - section.left_margin - section.right_margin - section.gutter
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run()
    run.add_picture(str(image_path), width=usable_width)


def render_md_into(doc: Document, md_path: Path) -> None:
    md = markdown.Markdown(extensions=["tables", "fenced_code", "sane_lists", "nl2br"])
    html = md.convert(md_path.read_text(encoding="utf-8"))
    soup = BeautifulSoup(html, "lxml")
    body = soup.find("body") or soup

    for element in body.children:
        if isinstance(element, NavigableString):
            if str(element).strip():
                p = doc.add_paragraph(str(element).strip())
                p.paragraph_format.first_line_indent = Cm(0.55)
            continue
        name = element.name
        if name in {"h1", "h2", "h3", "h4"}:
            add_heading(doc, element.get_text(" ", strip=True), int(name[1]))
        elif name == "p":
            img = element.find("img")
            if img is not None:
                add_image(doc, img, md_path.parent)
            else:
                add_paragraph(doc, element)
        elif name == "blockquote":
            add_blockquote(doc, element)
        elif name in {"ul", "ol"}:
            add_list(doc, element, ordered=(name == "ol"))
        elif name == "table":
            add_table(doc, element)
        elif name == "pre":
            add_code_block(doc, element)
        elif name == "hr":
            add_hr(doc)
        else:
            text = element.get_text(" ", strip=True)
            if text:
                p = doc.add_paragraph(text)
                p.paragraph_format.first_line_indent = Cm(0.55)


def new_doc(template_path: Path) -> Document:
    doc = Document(str(template_path))
    clear_body_keep_section(doc)
    setup_bookk_page(doc)
    setup_styles(doc)
    add_page_number_footer(doc)
    return doc


def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument("--template", default=DEFAULT_TEMPLATE, help="부크크 A5 워드서식 docx")
    ap.add_argument("--out", help="출력 docx (md 여러 개를 한 문서로, 파일마다 새 쪽)")
    ap.add_argument("--out-dir", help="출력 폴더 (md 하나당 docx 하나)")
    ap.add_argument("md_files", nargs="+")
    args = ap.parse_args()
    if bool(args.out) == bool(args.out_dir):
        ap.error("--out 또는 --out-dir 중 하나만 지정하세요.")

    template = Path(args.template)
    if not template.exists():
        ap.error(f"템플릿이 없습니다: {template}")

    if args.out:
        doc = new_doc(template)
        for idx, mf in enumerate(args.md_files):
            if idx > 0:
                doc.add_page_break()
            render_md_into(doc, Path(mf))
        out = Path(args.out)
        out.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(out))
        print("OK ->", out)
    else:
        out_dir = Path(args.out_dir)
        out_dir.mkdir(parents=True, exist_ok=True)
        for mf in args.md_files:
            mf = Path(mf)
            doc = new_doc(template)
            render_md_into(doc, mf)
            out = out_dir / (mf.stem + ".docx")
            doc.save(str(out))
            print("OK ->", out)


if __name__ == "__main__":
    main()
